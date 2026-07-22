"""Admin utilities for managing the Nritya.ai knowledge base."""

import json
import re
import shutil
import uuid
import zipfile
from datetime import datetime
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any

from PyPDF2 import PdfReader


BOOK_DIR = Path("books")
INDEX_DIR = Path("faiss_index")
BACKUP_DIR = Path("backups")
ADMIN_DATA_DIR = Path("admin_data")
BOOK_REGISTRY_FILE = ADMIN_DATA_DIR / "book_registry.json"
ANALYTICS_FILE = ADMIN_DATA_DIR / "analytics.json"
STATS_FILE = INDEX_DIR / "index_stats.json"

CATEGORIES = [
    "Prarambhik",
    "Madhyama",
    "Visharad",
    "Natyashastra",
    "Abhinaya",
    "General Theory",
]

ALLOWED_EXTENSIONS = {".pdf", ".txt", ".docx"}


def ensure_admin_dirs() -> None:
    BOOK_DIR.mkdir(exist_ok=True)
    INDEX_DIR.mkdir(exist_ok=True)
    BACKUP_DIR.mkdir(exist_ok=True)
    ADMIN_DATA_DIR.mkdir(exist_ok=True)


def read_json(path: Path, default: Any) -> Any:
    if not path.exists():
        return default
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return default


def write_json(path: Path, data: Any) -> None:
    path.parent.mkdir(exist_ok=True)
    path.write_text(json.dumps(data, indent=2), encoding="utf-8")


def slugify(text: str) -> str:
    cleaned = re.sub(r"[^a-zA-Z0-9]+", "_", text).strip("_").lower()
    return cleaned or f"book_{uuid.uuid4().hex[:8]}"


def infer_category(filename: str, requested_category: str | None = None) -> str:
    if requested_category in CATEGORIES:
        return requested_category
    lowered = filename.lower()
    for category in CATEGORIES:
        if category.lower().replace(" ", "") in lowered.replace("_", "").replace("-", ""):
            return category
    if "natya" in lowered:
        return "Natyashastra"
    if "abhinaya" in lowered:
        return "Abhinaya"
    return "General Theory"


def get_registry() -> list[dict[str, Any]]:
    ensure_admin_dirs()
    registry = read_json(BOOK_REGISTRY_FILE, [])
    existing_files = {item.get("filename") for item in registry}
    for path in BOOK_DIR.iterdir():
        if path.is_file() and path.suffix.lower() in ALLOWED_EXTENSIONS and path.name not in existing_files:
            registry.append(
                {
                    "id": slugify(path.stem),
                    "book_key": slugify(path.stem),
                    "name": path.name,
                    "filename": path.name,
                    "category": infer_category(path.name),
                    "upload_date": datetime.fromtimestamp(path.stat().st_mtime).isoformat(timespec="seconds"),
                    "status": "Indexed" if STATS_FILE.exists() else "Pending",
                    "pages": 0,
                    "chunks": 0,
                }
            )
    write_json(BOOK_REGISTRY_FILE, registry)
    return registry


def save_registry(registry: list[dict[str, Any]]) -> None:
    write_json(BOOK_REGISTRY_FILE, registry)


def infer_chapter_name(text: str, fallback: str) -> str:
    for line in text.splitlines()[:12]:
        cleaned = re.sub(r"\s+", " ", line).strip(" :-")
        if not cleaned:
            continue
        if re.search(r"\b(chapter|unit|lesson|adhyaya)\b", cleaned, re.I):
            return cleaned[:90]
        if cleaned.isupper() and 4 <= len(cleaned) <= 90:
            return cleaned.title()
    return fallback


def extract_pdf(path: Path) -> tuple[list[tuple[int, str]], int]:
    reader = PdfReader(str(path))
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append((index, text))
    return pages, len(reader.pages)


def extract_txt(path: Path) -> tuple[list[tuple[int, str]], int]:
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    return ([(1, text)] if text else []), 1


def extract_docx(path: Path) -> tuple[list[tuple[int, str]], int]:
    try:
        import docx

        document = docx.Document(str(path))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs).strip()
        return ([(1, text)] if text else []), 1
    except ImportError:
        with zipfile.ZipFile(path) as archive:
            xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
        text = re.sub(r"<[^>]+>", " ", xml)
        text = re.sub(r"\s+", " ", text).strip()
        return ([(1, text)] if text else []), 1


def extract_pages(path: Path) -> tuple[list[tuple[int, str]], int]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    if suffix == ".txt":
        return extract_txt(path)
    if suffix == ".docx":
        return extract_docx(path)
    raise ValueError(f"Unsupported file type: {suffix}")


def get_text_splitter():
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter
    except ImportError:
        from langchain.text_splitter import RecursiveCharacterTextSplitter
    return RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=80)


def build_documents_for_book(book: dict[str, Any], splitter) -> tuple[list[Any], int]:
    from langchain_core.documents import Document

    path = BOOK_DIR / book["filename"]
    pages, page_count = extract_pages(path)
    active_chapter = book.get("category") or "General Theory"
    documents = []
    for page_number, text in pages:
        active_chapter = infer_chapter_name(text, active_chapter)
        for chunk_number, chunk in enumerate(splitter.split_text(text), start=1):
            documents.append(
                Document(
                    page_content=chunk,
                    metadata={
                        "book_key": book["book_key"],
                        "source_book": book["filename"],
                        "book_label": book["name"],
                        "category": book.get("category", "General Theory"),
                        "chapter_name": active_chapter,
                        "page_number": page_number,
                        "chunk_number": chunk_number,
                    },
                )
            )
    return documents, page_count


def rebuild_faiss_index() -> dict[str, Any]:
    ensure_admin_dirs()
    registry = get_registry()
    from langchain_community.vectorstores import FAISS
    from langchain_huggingface import HuggingFaceEmbeddings

    splitter = get_text_splitter()
    documents = []
    book_stats = []

    for book in registry:
        path = BOOK_DIR / book.get("filename", "")
        if not path.exists():
            book["status"] = "Missing"
            continue
        try:
            book_documents, page_count = build_documents_for_book(book, splitter)
            book["pages"] = page_count
            book["chunks"] = len(book_documents)
            book["status"] = "Indexed"
            documents.extend(book_documents)
            book_stats.append(
                {
                    "key": book["book_key"],
                    "label": book["name"],
                    "filename": book["filename"],
                    "category": book.get("category", "General Theory"),
                    "pages": page_count,
                    "chunks": len(book_documents),
                    "upload_date": book.get("upload_date"),
                    "status": book["status"],
                }
            )
        except Exception as exc:
            book["status"] = f"Error: {exc}"

    if not documents:
        save_registry(registry)
        raise ValueError("No documents could be indexed.")

    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    db = FAISS.from_documents(documents, embeddings)
    db.save_local(str(INDEX_DIR))

    stats = {
        "books_indexed": len(book_stats),
        "total_pages_indexed": sum(book["pages"] for book in book_stats),
        "total_chunks_stored": sum(book["chunks"] for book in book_stats),
        "last_index_update": datetime.now().isoformat(timespec="seconds"),
        "books": book_stats,
    }
    write_json(STATS_FILE, stats)
    save_registry(registry)
    return stats


def register_upload(filename: str, content: bytes, category: str) -> dict[str, Any]:
    ensure_admin_dirs()
    suffix = Path(filename).suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise ValueError("Supported formats are PDF, TXT, and DOCX.")

    safe_name = re.sub(r"[^a-zA-Z0-9._ -]+", "", Path(filename).name).strip() or f"upload{suffix}"
    target = BOOK_DIR / safe_name
    if target.exists():
        target = BOOK_DIR / f"{target.stem}_{datetime.now().strftime('%Y%m%d%H%M%S')}{suffix}"
    target.write_bytes(content)

    registry = get_registry()
    book = {
        "id": uuid.uuid4().hex,
        "book_key": slugify(target.stem),
        "name": target.name,
        "filename": target.name,
        "category": infer_category(target.name, category),
        "upload_date": datetime.now().isoformat(timespec="seconds"),
        "status": "Uploaded",
        "pages": 0,
        "chunks": 0,
    }
    registry.append(book)
    save_registry(registry)
    return book


def delete_book(book_id: str) -> None:
    registry = get_registry()
    book = next((item for item in registry if item["id"] == book_id), None)
    if not book:
        raise ValueError("Book not found.")
    path = BOOK_DIR / book["filename"]
    if path.exists():
        path.unlink()
    save_registry([item for item in registry if item["id"] != book_id])


def update_book_category(book_id: str, category: str) -> dict[str, Any]:
    if category not in CATEGORIES:
        raise ValueError("Invalid category.")
    registry = get_registry()
    for book in registry:
        if book["id"] == book_id:
            book["category"] = category
            save_registry(registry)
            return book
    raise ValueError("Book not found.")


def get_admin_stats() -> dict[str, Any]:
    stats = read_json(STATS_FILE, {})
    analytics = get_analytics()
    return {
        "total_books_indexed": stats.get("books_indexed", 0),
        "total_pages_indexed": stats.get("total_pages_indexed", 0),
        "total_chunks": stats.get("total_chunks_stored", 0),
        "total_questions_asked": analytics.get("total_questions", 0),
        "most_searched_topics": top_items(analytics.get("topics", {}), 5),
        "last_database_update": stats.get("last_index_update", "Unknown"),
    }


def get_analytics() -> dict[str, Any]:
    return read_json(
        ANALYTICS_FILE,
        {
            "total_questions": 0,
            "questions": {},
            "topics": {},
            "books": {},
            "daily_usage": {},
        },
    )


def top_items(items: dict[str, int], limit: int = 8) -> list[dict[str, Any]]:
    return [
        {"label": label, "count": count}
        for label, count in sorted(items.items(), key=lambda item: item[1], reverse=True)[:limit]
    ]


def classify_topic(text: str) -> str:
    lowered = text.lower()
    topic_terms = {
        "Postures": ["aramandi", "murumandi", "posture", "mandi"],
        "Mudras": ["mudra", "hasta", "gesture"],
        "History": ["history", "origin", "temple", "revival"],
        "Devadasi Tradition": ["devadasi"],
        "Natyashastra": ["natyashastra", "bharata muni"],
        "Abhinaya": ["abhinaya", "bhava", "rasa"],
        "Theory Exams": ["exam", "marks", "viva"],
    }
    for topic, terms in topic_terms.items():
        if any(term in lowered for term in terms):
            return topic
    return "General Theory"


def record_question(question: str, retrieval: dict[str, Any] | None = None) -> None:
    analytics = get_analytics()
    today = datetime.now().date().isoformat()
    normalized_question = question.strip()[:160]
    topic = classify_topic(question)
    analytics["total_questions"] = analytics.get("total_questions", 0) + 1
    analytics.setdefault("questions", {})[normalized_question] = analytics.setdefault("questions", {}).get(normalized_question, 0) + 1
    analytics.setdefault("topics", {})[topic] = analytics.setdefault("topics", {}).get(topic, 0) + 1
    analytics.setdefault("daily_usage", {})[today] = analytics.setdefault("daily_usage", {}).get(today, 0) + 1
    if retrieval:
        for book in retrieval.get("books_used", []):
            analytics.setdefault("books", {})[book] = analytics.setdefault("books", {}).get(book, 0) + 1
    write_json(ANALYTICS_FILE, analytics)


def get_analytics_dashboard() -> dict[str, Any]:
    analytics = get_analytics()
    return {
        "most_asked_questions": top_items(analytics.get("questions", {}), 8),
        "frequently_studied_topics": top_items(analytics.get("topics", {}), 8),
        "popular_books": top_items(analytics.get("books", {}), 8),
        "daily_usage_trends": [
            {"date": date, "count": count}
            for date, count in sorted(analytics.get("daily_usage", {}).items())[-14:]
        ],
    }


def create_backup() -> Path:
    ensure_admin_dirs()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    backup_path = BACKUP_DIR / f"nritya_backup_{timestamp}"
    if backup_path.exists() or backup_path.with_suffix(".zip").exists():
        backup_path = BACKUP_DIR / f"nritya_backup_{timestamp}_{uuid.uuid4().hex[:6]}"
    backup_path.mkdir()
    if INDEX_DIR.exists():
        shutil.copytree(INDEX_DIR, backup_path / "faiss_index")
    if BOOK_DIR.exists():
        shutil.copytree(BOOK_DIR, backup_path / "books")
    if ADMIN_DATA_DIR.exists():
        shutil.copytree(ADMIN_DATA_DIR, backup_path / "admin_data")
    archive = shutil.make_archive(str(backup_path), "zip", backup_path)
    shutil.rmtree(backup_path)
    return Path(archive)


def _validate_backup_archive(archive_path: Path) -> None:
    with zipfile.ZipFile(archive_path) as archive:
        for member in archive.infolist():
            member_path = Path(member.filename)
            if member_path.is_absolute() or ".." in member_path.parts:
                raise ValueError("Backup archive contains unsafe paths.")


def _find_restore_root(extracted_dir: Path) -> Path:
    expected = {"faiss_index", "books", "admin_data"}
    if any((extracted_dir / name).exists() for name in expected):
        return extracted_dir

    children = [child for child in extracted_dir.iterdir() if child.is_dir()]
    for child in children:
        if any((child / name).exists() for name in expected):
            return child

    raise ValueError("Backup archive does not contain a Nritya.ai knowledge base.")


def restore_backup(archive_name: str, content: bytes) -> dict[str, Any]:
    ensure_admin_dirs()
    if Path(archive_name).suffix.lower() != ".zip":
        raise ValueError("Only .zip backup archives can be imported.")

    incoming = BACKUP_DIR / f"restore_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{Path(archive_name).name}"
    incoming.write_bytes(content)
    _validate_backup_archive(incoming)

    safety_backup = create_backup()
    with TemporaryDirectory(dir=BACKUP_DIR) as temp_dir:
        temp_path = Path(temp_dir)
        with zipfile.ZipFile(incoming) as archive:
            archive.extractall(temp_path)

        restore_root = _find_restore_root(temp_path)
        for directory_name in ("faiss_index", "books", "admin_data"):
            source = restore_root / directory_name
            target = Path(directory_name)
            if not source.exists():
                continue
            if target.exists():
                shutil.rmtree(target)
            shutil.copytree(source, target)

    return {
        "status": "restored",
        "imported_backup": incoming.name,
        "safety_backup": safety_backup.name,
    }
